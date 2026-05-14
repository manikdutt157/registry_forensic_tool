import csv
import html
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple


REPORT_ARTIFACTS = (
    ("System Info", "System_Info", "system_info.csv"),
    ("Persistence", "Persistence", "run_keys.csv"),
    ("USB History", "USB_History", "usb_devices.csv"),
    ("User Activity", "User_Activity", "recent_activity.csv"),
    ("Execution History", "Execution_History", "execution_history.csv"),
    ("Network Info", "Network_Info", "network_profiles.csv"),
)

REPORT_TITLE = "Windows Registry Forensic Examination Report"
REPORT_PURPOSE = (
    "This report summarizes registry artifacts extracted and parsed by the "
    "Windows Registry Forensic Tool. Findings should be validated against the "
    "original evidence and supporting forensic sources before final case conclusions."
)
METHODOLOGY_STEPS = (
    "Acquire registry hives from the examined Windows system.",
    "Parse selected registry locations for system, user activity, execution, persistence, USB, and network artifacts.",
    "Preserve parsed results as CSV evidence tables in the case output directory.",
    "Generate this report from the parsed CSV evidence for review and documentation.",
)
LIMITATIONS = (
    "Registry timestamps and values can be affected by system activity, anti-forensic activity, corruption, or parser limits.",
    "This report is a structured summary of parsed artifacts and does not replace full disk, memory, or timeline analysis.",
    "Rows may be truncated in document formats when evidence volume is high; CSV files remain the complete parsed output.",
)


def read_csv(csv_path):
    with open(csv_path, "r", encoding="utf-8", newline="") as file:
        reader = csv.reader(file)
        headers = next(reader, [])
        rows = list(reader)
    return headers, rows


def collect_report_data(
    case_dir: Path, artifacts: Sequence[Tuple[str, str, str]] = REPORT_ARTIFACTS
) -> Tuple[List[Tuple[str, int, str]], List[Dict[str, object]]]:
    summary_rows: List[Tuple[str, int, str]] = []
    sections: List[Dict[str, object]] = []
    for title, folder, filename in artifacts:
        csv_path = case_dir / folder / filename
        if not csv_path.exists():
            summary_rows.append((title, 0, "Missing"))
            sections.append(
                {
                    "title": title,
                    "status": "missing",
                    "source": str(csv_path),
                    "headers": [],
                    "rows": [],
                }
            )
            continue

        headers, rows = read_csv(csv_path)
        summary_rows.append((title, len(rows), filename))
        sections.append(
            {
                "title": title,
                "status": "ok" if headers else "empty",
                "source": filename,
                "headers": headers,
                "rows": rows,
            }
        )
    return summary_rows, sections


def generate_html_report(case_dir, artifacts=REPORT_ARTIFACTS):
    case_dir = Path(case_dir)
    report_dir = case_dir / "Reports"
    report_dir.mkdir(parents=True, exist_ok=True)
    report_path = report_dir / f"registry_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"

    summary_rows, report_sections = collect_report_data(case_dir, artifacts)
    sections = []
    for section in report_sections:
        title = str(section["title"])
        status = str(section["status"])
        if status == "missing":
            sections.append(render_missing_section(title, Path(str(section["source"]))))
        else:
            sections.append(
                render_table_section(
                    title,
                    str(section["source"]),
                    list(section["headers"]),  # type: ignore[arg-type]
                    list(section["rows"]),  # type: ignore[arg-type]
                )
            )

    report_path.write_text(render_report(case_dir, summary_rows, sections), encoding="utf-8")
    return report_path


def render_report(case_dir, summary_rows, sections):
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    case_name = html.escape(case_dir.name)
    summary = "\n".join(
        f"<tr><td>{html.escape(title)}</td><td>{count}</td><td>{html.escape(source)}</td></tr>"
        for title, count, source in summary_rows
    )
    body_sections = "\n".join(sections)

    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <title>{REPORT_TITLE} - {case_name}</title>
  <style>
    :root {{
      --bg: #f5f7fb;
      --panel: #ffffff;
      --ink: #17202a;
      --muted: #5f6b7a;
      --line: #dce3ed;
      --accent: #1f6feb;
      --head: #eef4ff;
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      background: var(--bg);
      color: var(--ink);
      font-family: "Segoe UI", Arial, sans-serif;
      line-height: 1.45;
    }}
    header {{
      background: #111827;
      color: #fff;
      padding: 28px 36px;
    }}
    header h1 {{ margin: 0 0 8px; font-size: 28px; }}
    header p {{ margin: 0; color: #cbd5e1; }}
    main {{ padding: 24px 36px 42px; }}
    section {{
      background: var(--panel);
      border: 1px solid var(--line);
      border-radius: 8px;
      margin-bottom: 22px;
      overflow: hidden;
      box-shadow: 0 1px 2px rgba(16, 24, 40, 0.05);
    }}
    section h2 {{
      margin: 0;
      padding: 16px 18px;
      border-bottom: 1px solid var(--line);
      font-size: 18px;
    }}
    .meta {{
      padding: 0 18px 14px;
      color: var(--muted);
      font-size: 13px;
    }}
    .table-wrap {{ overflow-x: auto; }}
    table {{
      width: 100%;
      border-collapse: collapse;
      font-size: 13px;
    }}
    th, td {{
      border-bottom: 1px solid var(--line);
      padding: 8px 10px;
      text-align: left;
      vertical-align: top;
      max-width: 520px;
      overflow-wrap: anywhere;
    }}
    th {{
      background: var(--head);
      color: #24364b;
      position: sticky;
      top: 0;
    }}
    tr:nth-child(even) td {{ background: #fafcff; }}
    .summary td:first-child {{ font-weight: 600; }}
    .field-grid {{
      display: grid;
      grid-template-columns: 180px 1fr;
      gap: 8px 16px;
      padding: 18px;
    }}
    .field-label {{ color: var(--muted); font-weight: 700; }}
    ul {{ margin: 0; padding: 0 18px 18px 38px; }}
    .empty {{ padding: 18px; color: var(--muted); }}
    footer {{
      color: var(--muted);
      font-size: 12px;
      padding-top: 6px;
    }}
    @media print {{
      body {{ background: #fff; }}
      header {{ background: #fff; color: #000; border-bottom: 2px solid #000; }}
      header p, footer {{ color: #333; }}
      section {{ box-shadow: none; break-inside: avoid; }}
      th {{ position: static; }}
    }}
  </style>
</head>
<body>
  <header>
    <h1>{REPORT_TITLE}</h1>
    <p>Case: {case_name} | Generated: {html.escape(generated_at)}</p>
  </header>
  <main>
    <section>
      <h2>Case Overview</h2>
      <div class="field-grid">
        <div class="field-label">Case Folder</div><div>{case_name}</div>
        <div class="field-label">Report Generated</div><div>{html.escape(generated_at)}</div>
        <div class="field-label">Tool</div><div>Windows Registry Forensic Tool</div>
        <div class="field-label">Report Type</div><div>Registry artifact examination summary</div>
      </div>
    </section>
    <section>
      <h2>Executive Summary</h2>
      <div class="empty">{html.escape(REPORT_PURPOSE)}</div>
    </section>
    <section>
      <h2>Methodology</h2>
      <ul>{"".join(f"<li>{html.escape(step)}</li>" for step in METHODOLOGY_STEPS)}</ul>
    </section>
    <section>
      <h2>Findings Summary</h2>
      <div class="table-wrap">
        <table class="summary">
          <thead><tr><th>Category</th><th>Rows</th><th>Source</th></tr></thead>
          <tbody>{summary}</tbody>
        </table>
      </div>
    </section>
    {body_sections}
    <section>
      <h2>Limitations</h2>
      <ul>{"".join(f"<li>{html.escape(item)}</li>" for item in LIMITATIONS)}</ul>
    </section>
    <footer>Generated by Windows Registry Forensic Tool.</footer>
  </main>
</body>
</html>
"""


def render_missing_section(title, csv_path):
    return f"""<section>
  <h2>{html.escape(title)}</h2>
  <div class="empty">Missing source file: {html.escape(str(csv_path))}</div>
</section>"""


def render_table_section(title, filename, headers, rows):
    safe_title = html.escape(title)
    safe_filename = html.escape(filename)
    if not headers:
        return f"""<section>
  <h2>{safe_title}</h2>
  <div class="empty">No columns were found in {safe_filename}.</div>
</section>"""

    header_html = "".join(f"<th>{html.escape(header.replace('_', ' ').title())}</th>" for header in headers)
    row_html = []
    for row in rows:
        padded = row + [""] * (len(headers) - len(row))
        cells = "".join(f"<td>{html.escape(str(value))}</td>" for value in padded[: len(headers)])
        row_html.append(f"<tr>{cells}</tr>")

    body = "\n".join(row_html) if row_html else f"<tr><td colspan=\"{len(headers)}\">No rows found.</td></tr>"
    return f"""<section>
  <h2>{safe_title}</h2>
  <div class="meta">{len(rows)} rows from {safe_filename}</div>
  <div class="table-wrap">
    <table>
      <thead><tr>{header_html}</tr></thead>
      <tbody>{body}</tbody>
    </table>
  </div>
</section>"""


def generate_pdf_report(case_dir, artifacts=REPORT_ARTIFACTS):
    """
    Generate a PDF report under <case_dir>/Reports using reportlab.
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    case_dir = Path(case_dir)
    report_dir = case_dir / "Reports"
    report_dir.mkdir(parents=True, exist_ok=True)
    report_path = report_dir / f"registry_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"

    summary_rows, report_sections = collect_report_data(case_dir, artifacts)

    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    h2_style = styles["Heading2"]
    meta_style = ParagraphStyle("Meta", parent=styles["Normal"], textColor=colors.HexColor("#4b5563"), fontSize=9)
    cell_style = ParagraphStyle("Cell", parent=styles["Normal"], fontSize=8, leading=10)

    doc = SimpleDocTemplate(
        str(report_path),
        pagesize=A4,
        leftMargin=1.6 * cm,
        rightMargin=1.6 * cm,
        topMargin=1.4 * cm,
        bottomMargin=1.4 * cm,
        title=f"{REPORT_TITLE} - {case_dir.name}",
    )

    story = []
    story.append(Paragraph(REPORT_TITLE, title_style))
    story.append(Paragraph(f"Case: {html.escape(case_dir.name)}", meta_style))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", meta_style))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Case Overview", h2_style))
    overview_data = [
        ["Case Folder", case_dir.name],
        ["Report Generated", datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
        ["Tool", "Windows Registry Forensic Tool"],
        ["Report Type", "Registry artifact examination summary"],
    ]
    overview_table = Table(overview_data, hAlign="LEFT", colWidths=[doc.width * 0.28, doc.width * 0.72])
    overview_table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#dce3ed")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(overview_table)
    story.append(Spacer(1, 10))

    story.append(Paragraph("Executive Summary", h2_style))
    story.append(Paragraph(html.escape(REPORT_PURPOSE), meta_style))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Methodology", h2_style))
    for step in METHODOLOGY_STEPS:
        story.append(Paragraph(f"- {html.escape(step)}", meta_style))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Findings Summary", h2_style))
    summary_data = [["Category", "Rows", "Source"]] + [[t, str(c), s] for (t, c, s) in summary_rows]
    summary_table = Table(
        summary_data,
        repeatRows=1,
        hAlign="LEFT",
        colWidths=[doc.width * 0.45, doc.width * 0.12, doc.width * 0.43],
    )
    summary_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef4ff")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#24364b")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#dce3ed")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 2),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#fafcff")]),
            ]
        )
    )
    story.append(summary_table)
    story.append(PageBreak())

    for section in report_sections:
        title = str(section["title"])
        status = str(section["status"])
        story.append(Paragraph(title, h2_style))
        if status == "missing":
            story.append(Paragraph(f"Missing source file: {html.escape(str(section['source']))}", meta_style))
            story.append(Spacer(1, 10))
            continue

        headers: List[str] = list(section["headers"])  # type: ignore[assignment]
        rows: List[List[str]] = [list(r) for r in section["rows"]]  # type: ignore[assignment]
        story.append(Paragraph(f"{len(rows)} rows from {html.escape(str(section['source']))}", meta_style))
        story.append(Spacer(1, 6))

        if not headers:
            story.append(Paragraph("No columns were found in this CSV.", meta_style))
            story.append(Spacer(1, 10))
            continue

        # Build a table with wrapped cells to avoid overflow.
        table_data: List[List[object]] = [[Paragraph(html.escape(h.replace("_", " ").title()), cell_style) for h in headers]]
        max_rows = 500  # keep PDFs sane; still include a lot of evidence
        truncated = rows[:max_rows]
        for row in truncated:
            padded = row + [""] * (len(headers) - len(row))
            table_data.append([Paragraph(html.escape(str(v)), cell_style) for v in padded[: len(headers)]])
        if len(rows) > max_rows:
            table_data.append(
                [
                    Paragraph(
                        f"(Truncated) Showing {max_rows} of {len(rows)} rows. See CSV for full data.",
                        meta_style,
                    )
                ]
                + [Paragraph("", meta_style) for _ in range(len(headers) - 1)]
            )

        tbl = Table(table_data, repeatRows=1, hAlign="LEFT")
        tbl.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef4ff")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#24364b")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#dce3ed")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 1),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 1),
                    ("TOPPADDING", (0, 0), (-1, -1), 1),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#fafcff")]),
                ]
            )
        )
        # Force the table to fit the page width even with many columns.
        if headers:
            col_width = max(6, int(doc.width / len(headers)))
            tbl._argW = [col_width for _ in range(len(headers))]
        story.append(tbl)
        story.append(PageBreak())

    story.append(Paragraph("Limitations", h2_style))
    for item in LIMITATIONS:
        story.append(Paragraph(f"- {html.escape(item)}", meta_style))

    # Remove the last PageBreak if present (cosmetic).
    if story and story[-1].__class__.__name__ == "PageBreak":
        story.pop()
    doc.build(story)
    return report_path


def generate_docx_report(case_dir, artifacts=REPORT_ARTIFACTS):
    """
    Generate a DOCX report under <case_dir>/Reports using python-docx.
    """
    from docx import Document

    case_dir = Path(case_dir)
    report_dir = case_dir / "Reports"
    report_dir.mkdir(parents=True, exist_ok=True)
    report_path = report_dir / f"registry_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    summary_rows, report_sections = collect_report_data(case_dir, artifacts)

    doc = Document()
    doc.add_heading(REPORT_TITLE, level=0)
    doc.add_paragraph(f"Case: {case_dir.name}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("Case Overview", level=1)
    overview_table = doc.add_table(rows=4, cols=2)
    overview_rows = [
        ("Case Folder", case_dir.name),
        ("Report Generated", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("Tool", "Windows Registry Forensic Tool"),
        ("Report Type", "Registry artifact examination summary"),
    ]
    for index, (label, value) in enumerate(overview_rows):
        overview_table.rows[index].cells[0].text = label
        overview_table.rows[index].cells[1].text = value

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(REPORT_PURPOSE)

    doc.add_heading("Methodology", level=1)
    for step in METHODOLOGY_STEPS:
        doc.add_paragraph(step, style="List Bullet")

    doc.add_heading("Findings Summary", level=1)
    summary_table = doc.add_table(rows=1, cols=3)
    hdr = summary_table.rows[0].cells
    hdr[0].text = "Category"
    hdr[1].text = "Rows"
    hdr[2].text = "Source"
    for title, count, source in summary_rows:
        row_cells = summary_table.add_row().cells
        row_cells[0].text = str(title)
        row_cells[1].text = str(count)
        row_cells[2].text = str(source)

    for section in report_sections:
        title = str(section["title"])
        status = str(section["status"])
        doc.add_page_break()
        doc.add_heading(title, level=1)
        if status == "missing":
            doc.add_paragraph(f"Missing source file: {section['source']}")
            continue

        headers: List[str] = list(section["headers"])  # type: ignore[assignment]
        rows: List[List[str]] = [list(r) for r in section["rows"]]  # type: ignore[assignment]
        doc.add_paragraph(f"{len(rows)} rows from {section['source']}")

        if not headers:
            doc.add_paragraph("No columns were found in this CSV.")
            continue

        # DOCX can get huge; cap rows and note truncation.
        max_rows = 2000
        shown = rows[:max_rows]

        tbl = doc.add_table(rows=1, cols=len(headers))
        for idx, h in enumerate(headers):
            tbl.rows[0].cells[idx].text = h.replace("_", " ").title()
        for r in shown:
            padded = r + [""] * (len(headers) - len(r))
            cells = tbl.add_row().cells
            for idx, v in enumerate(padded[: len(headers)]):
                cells[idx].text = str(v)

        if len(rows) > max_rows:
            doc.add_paragraph(f"(Truncated) Showing {max_rows} of {len(rows)} rows. See CSV for full data.")

    doc.add_heading("Limitations", level=1)
    for item in LIMITATIONS:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(str(report_path))
    return report_path


def generate_reports(case_dir, formats: Iterable[str] = ("pdf", "docx")) -> Dict[str, Path]:
    """
    Convenience wrapper: generate multiple report formats.
    """
    outputs: Dict[str, Path] = {}
    for fmt in formats:
        f = fmt.strip().lower()
        if f == "html":
            outputs[f] = Path(generate_html_report(case_dir))
        elif f == "pdf":
            outputs[f] = Path(generate_pdf_report(case_dir))
        elif f in ("doc", "docx"):
            outputs["docx"] = Path(generate_docx_report(case_dir))
        else:
            raise ValueError(f"Unsupported report format: {fmt}")
    return outputs
